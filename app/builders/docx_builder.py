"""
Output builders for document layout reconstruction.
"""
from docx import Document as WordDocument
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap

from app.models import (
    Document, Page, TextObject, TableObject, 
    ImageObject, PageElement, SeparatorObject
)

class DocxBuilder:
    def build(self, doc_model: Document, path: str):
        print(f"DocxBuilder: Building document at {path}...")
        
        word_doc = WordDocument()
        
        # Set margins to zero to allow full page canvas usage
        section = word_doc.sections[0]
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        
        for page in doc_model.pages:
            if page.pageNumber > 1:
                word_doc.add_page_break()
                
            # Calculate scaling factors
            # Word uses Inches or EMU. 
            # We have pixels.
            # Let's assume the page width in Word is standard Letter (8.5 inches)
            # Scale factor = 8.5 inches / image_width_pixels
            
            target_width_inches = 8.5
            scale_factor = target_width_inches / page.width
            
            for element in page.elements:
                self._process_element(word_doc, element, scale_factor)
                
        word_doc.save(path)
        print("DocxBuilder: Save complete.")

    def _process_element(self, word_doc, element: PageElement, scale: float):
        # Use Frames for exact positioning (w:framePr)
        # This works for both Text and Images (if image is inline in a frame)
        
        if isinstance(element, TextObject):
            self._create_frame(word_doc, element, scale, is_image=False)
        elif isinstance(element, ImageObject):
            self._create_frame(word_doc, element, scale, is_image=True)
        elif isinstance(element, TableObject):
            self._add_table(word_doc, element, scale)
        elif isinstance(element, SeparatorObject):
            self._add_separator(word_doc, element, scale)

    def _create_frame(self, word_doc, element: PageElement, scale: float, is_image: bool):
        # Create a paragraph
        p = word_doc.add_paragraph()
        
        # Calculate dimensions in Twips (1/1440 inch)
        # element.bbox is in pixels. scale converts to inches.
        # Twips = Inches * 1440
        
        x_twips = int(element.bbox.x * scale * 1440)
        y_twips = int(element.bbox.y * scale * 1440)
        w_twips = int(element.bbox.width * scale * 1440)
        h_twips = int(element.bbox.height * scale * 1440)
        
        # Add frame properties to the paragraph
        # <w:framePr w:w="W" w:h="H" w:x="X" w:y="Y" w:hRule="exact" w:wrap="around" w:vAnchor="page" w:hAnchor="page"/>
        
        p_pr = p._p.get_or_add_pPr()
        frame_pr = OxmlElement('w:framePr')
        frame_pr.set(qn('w:w'), str(w_twips))
        frame_pr.set(qn('w:h'), str(h_twips))
        frame_pr.set(qn('w:x'), str(x_twips))
        frame_pr.set(qn('w:y'), str(y_twips))
        frame_pr.set(qn('w:hRule'), 'exact')
        frame_pr.set(qn('w:wrap'), 'around') # or 'auto'
        frame_pr.set(qn('w:vAnchor'), 'page')
        frame_pr.set(qn('w:hAnchor'), 'page')
        
        p_pr.append(frame_pr)
        
        # Add content to the paragraph
        if is_image:
            run = p.add_run()
            try:
                # Add picture inline - the frame positions the paragraph
                run.add_picture(element.imagePath, width=Inches(element.bbox.width * scale))
            except Exception as e:
                print(f"Error adding image {element.imagePath}: {e}")
        else:
            # Text content
            run = p.add_run(element.rawText)
            
            # Styling
            if element.fontName:
                run.font.name = element.fontName
            if element.fontSize:
                run.font.size = Pt(element.fontSize)
            if element.isBold:
                run.font.bold = True
            if element.isItalic:
                run.font.italic = True
                

            if element.fontColor:
                try:
                    # element.fontColor is hex string "RRGGBB"
                    r = int(element.fontColor[0:2], 16)
                    g = int(element.fontColor[2:4], 16)
                    b = int(element.fontColor[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
                except Exception:
                    pass
            
            # Alignment
            if element.alignment == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif element.alignment == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif element.alignment == "justify":
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_table(self, word_doc, table_obj: TableObject, scale: float):
        """
        Add a properly formatted and positioned table to the document.
        Uses absolute positioning via tblpPr to place table at exact coordinates.
        """
        if table_obj.rowCount == 0 or table_obj.colCount == 0:
            return
            
        # Create table with correct dimensions
        table = word_doc.add_table(rows=table_obj.rowCount, cols=table_obj.colCount)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Calculate dimensions in Twips (1/1440 inch)
        x_twips = int(table_obj.bbox.x * scale * 1440)
        y_twips = int(table_obj.bbox.y * scale * 1440)
        table_width_twips = int(table_obj.bbox.width * scale * 1440)
        table_height_twips = int(table_obj.bbox.height * scale * 1440)
        
        # Add table properties for proper layout
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        
        # Table positioning properties - anchor to page
        tblpPr = OxmlElement('w:tblpPr')
        tblpPr.set(qn('w:leftFromText'), '0')
        tblpPr.set(qn('w:rightFromText'), '0')
        tblpPr.set(qn('w:topFromText'), '0')
        tblpPr.set(qn('w:bottomFromText'), '0')
        tblpPr.set(qn('w:vertAnchor'), 'page')
        tblpPr.set(qn('w:horzAnchor'), 'page')
        tblpPr.set(qn('w:tblpX'), str(x_twips))
        tblpPr.set(qn('w:tblpY'), str(y_twips))
        tblPr.insert(0, tblpPr)
        
        # Table overlap - prevent overlap with other elements
        tblOverlap = OxmlElement('w:tblOverlap')
        tblOverlap.set(qn('w:val'), 'never')
        tblPr.append(tblOverlap)
        
        # Fixed table layout - prevents column resizing
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        
        # Set table width exactly
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(table_width_twips))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        
        # Ensure tblPr is attached
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
        
        # Calculate column width - divide table width evenly
        table_width_inches = table_obj.bbox.width * scale
        col_width = Inches(table_width_inches / max(1, table_obj.colCount))

        
        for row_idx, row_obj in enumerate(table_obj.rows):
            if row_idx >= len(table.rows):
                continue
                
            row = table.rows[row_idx]
            
            # Set row height if we have bbox info
            if row_obj.cells:
                first_cell = row_obj.cells[0]
                if first_cell.bbox:
                    row_height = first_cell.bbox.height * scale
                    row.height = Inches(row_height)
            
            for col_idx, cell_obj in enumerate(row_obj.cells):
                if col_idx >= len(row.cells):
                    continue
                    
                cell = row.cells[col_idx]
                
                # Set cell width
                cell.width = col_width
                
                # Clear default paragraph if we're adding content
                if cell_obj.content:
                    # Remove default empty paragraph
                    if cell.paragraphs:
                        cell.paragraphs[0].clear()
                    
                    for content_item in cell_obj.content:
                        if isinstance(content_item, TextObject):
                            # Use existing paragraph or create new one
                            if cell.paragraphs and not cell.paragraphs[0].text:
                                p = cell.paragraphs[0]
                            else:
                                p = cell.add_paragraph()
                            
                            run = p.add_run(content_item.rawText)
                            
                            # Apply text styling
                            if content_item.fontSize and content_item.fontSize > 0:
                                run.font.size = Pt(content_item.fontSize)
                            else:
                                run.font.size = Pt(9)  # Default cell font size
                            
                            if content_item.fontName:
                                run.font.name = content_item.fontName
                            
                            if content_item.isBold:
                                run.font.bold = True
                            
                            if content_item.isItalic:
                                run.font.italic = True
                            
                            # Apply color
                            if content_item.fontColor:
                                try:
                                    r = int(content_item.fontColor[0:2], 16)
                                    g = int(content_item.fontColor[2:4], 16)
                                    b = int(content_item.fontColor[4:6], 16)
                                    run.font.color.rgb = RGBColor(r, g, b)
                                except:
                                    pass
                            
                            # Cell alignment - center by default for tables
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    def _add_separator(self, word_doc, element: SeparatorObject, scale: float):
        # Use a frame to position the separator
        # Inside the frame, add a paragraph with a bottom border
        
        p = word_doc.add_paragraph()
        
        # Calculate dimensions
        x_twips = int(element.bbox.x * scale * 1440)
        y_twips = int(element.bbox.y * scale * 1440)
        w_twips = int(element.bbox.width * scale * 1440)
        h_twips = int(max(element.thickness * scale * 1440, 20)) # Min height for visibility
        
        # Frame properties
        p_pr = p._p.get_or_add_pPr()
        frame_pr = OxmlElement('w:framePr')
        frame_pr.set(qn('w:w'), str(w_twips))
        frame_pr.set(qn('w:h'), str(h_twips))
        frame_pr.set(qn('w:x'), str(x_twips))
        frame_pr.set(qn('w:y'), str(y_twips))
        frame_pr.set(qn('w:hRule'), 'exact')
        frame_pr.set(qn('w:wrap'), 'around')
        frame_pr.set(qn('w:vAnchor'), 'page')
        frame_pr.set(qn('w:hAnchor'), 'page')
        p_pr.append(frame_pr)
        
        # Add border to paragraph to simulate line
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6') # 1/8 pt units. 6 = 0.75pt
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), element.color)
        pbdr.append(bottom)
        p_pr.append(pbdr)


class HtmlBuilder:
    def build(self, doc_model: Document, path: str):
        print(f"HtmlBuilder: Building document at {path}...")
        
        html_content = ["<!DOCTYPE html>", "<html>", "<head>", "<style>"]
        
        # CSS for exact layout
        html_content.append("""
            body { margin: 0; padding: 20px; background-color: #f0f0f0; }
            .page { 
                position: relative; 
                background-color: white; 
                box-shadow: 0 0 10px rgba(0,0,0,0.1); 
                margin-bottom: 20px;
                overflow: hidden; /* Ensure content doesn't spill */
            }
            .element { position: absolute; overflow: hidden; }
            .text-element { font-family: Arial, sans-serif; line-height: 1.2; }
            .image-element img { width: 100%; height: 100%; object-fit: contain; }
            .table-element { border-collapse: collapse; background-color: white; }
            .table-element td { border: 1px solid black; padding: 2px; vertical-align: top; }
        """)
        
        html_content.extend(["</style>", "</head>", "<body>"])
        
        for page in doc_model.pages:
            # Page container
            # We use the pixel dimensions directly
            html_content.append(f'<div class="page" style="width: {page.width}px; height: {page.height}px;">')
            
            for element in page.elements:
                self._process_element(html_content, element)
                
            html_content.append("</div>") # End page
            
        html_content.extend(["</body>", "</html>"])
        
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(html_content))
            
        print("HtmlBuilder: Save complete.")

    def _process_element(self, html_content, element: PageElement):
        # Common styles for position
        style = f"left: {element.bbox.x}px; top: {element.bbox.y}px; width: {element.bbox.width}px; height: {element.bbox.height}px;"
        
        if isinstance(element, TextObject):
            # Font styling
            font_style = f"font-size: {element.fontSize}pt;"
            if element.isBold: font_style += " font-weight: bold;"
            if element.isItalic: font_style += " font-style: italic;"
            if element.fontColor: font_style += f" color: #{element.fontColor};"
            
            # Alignment
            align_style = f"text-align: {element.alignment};"
            
            html_content.append(f'<div class="element text-element" style="{style} {font_style} {align_style}">')
            html_content.append(element.rawText.replace("\n", "<br>"))
            html_content.append("</div>")
            
        elif isinstance(element, ImageObject):
            html_content.append(f'<div class="element image-element" style="{style}">')
            html_content.append(f'<img src="{element.imagePath}" alt="Detected Image">')
            html_content.append("</div>")
            
        elif isinstance(element, TableObject):
            # For tables, we might need to be careful. 
            # If we just put a table in a div of fixed size, it might overflow or be too small.
            # We'll try to fit it.
            html_content.append(f'<div class="element table-element-container" style="{style} overflow: auto;">')
            html_content.append('<table class="table-element" style="width: 100%; height: 100%;">')
            
            for row in element.rows:
                html_content.append("<tr>")
                for cell in row.cells:
                    # Cell content
                    cell_html = ""
                    for item in cell.content:
                        if isinstance(item, TextObject):
                            cell_html += f"<div>{item.rawText}</div>"
                    
                    colspan = f' colspan="{cell.colSpan}"' if cell.colSpan > 1 else ""
                    rowspan = f' rowspan="{cell.rowSpan}"' if cell.rowSpan > 1 else ""
                    
                    html_content.append(f'<td{colspan}{rowspan}>{cell_html}</td>')
                html_content.append("</tr>")
            
            html_content.append("</table>")
            html_content.append("</div>")
            
        elif isinstance(element, SeparatorObject):
            # Render as a div with a bottom border
            sep_style = f"border-bottom: {max(1, element.thickness)}px solid #{element.color};"
            html_content.append(f'<div class="element separator-element" style="{style} {sep_style}"></div>')
class DocBookBuilder:
    def build(self, doc_model: Document, path: str):
        print(f"DocBookBuilder: Building document at {path}...")
        
        # We will build the XML string manually to ensure control over namespaces and formatting
        # DocBook 5.0
        
        xml_lines = []
        xml_lines.append('<?xml version="1.0" encoding="UTF-8"?>')
        xml_lines.append('<article xmlns="http://docbook.org/ns/docbook"')
        xml_lines.append('         xmlns:xlink="http://www.w3.org/1999/xlink"')
        xml_lines.append('         xmlns:layout="http://example.com/layout"') # Custom namespace for exact layout
        xml_lines.append('         version="5.0">')
        
        xml_lines.append(f'  <info><title>Reconstructed Document {doc_model.docID}</title></info>')
        
        for page in doc_model.pages:
            xml_lines.append(f'  <section role="page" layout:width="{page.width}" layout:height="{page.height}" label="{page.pageNumber}">')
            xml_lines.append(f'    <title>Page {page.pageNumber}</title>')
            
            for element in page.elements:
                self._process_element(xml_lines, element)
                
            xml_lines.append('  </section>')
            
        xml_lines.append('</article>')
        
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(xml_lines))
            
        print("DocBookBuilder: Save complete.")

    def _process_element(self, xml_lines, element: PageElement):
        # Layout attributes
        layout_attrs = f'layout:x="{element.bbox.x}" layout:y="{element.bbox.y}" layout:width="{element.bbox.width}" layout:height="{element.bbox.height}"'
        
        if isinstance(element, TextObject):
            # Clean text for XML
            safe_text = element.rawText.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            
            # Attributes for styling
            role = "text"
            if element.isBold: role += " bold"
            if element.isItalic: role += " italic"
            
            xml_lines.append(f'    <para role="{role}" {layout_attrs}>{safe_text}</para>')
            
        elif isinstance(element, ImageObject):
            xml_lines.append(f'    <mediaobject {layout_attrs}>')
            xml_lines.append('      <imageobject>')
            xml_lines.append(f'        <imagedata fileref="{element.imagePath}" contentwidth="{element.bbox.width}px" contentdepth="{element.bbox.height}px"/>')
            xml_lines.append('      </imageobject>')
            xml_lines.append('    </mediaobject>')
            
        elif isinstance(element, TableObject):
            xml_lines.append(f'    <informaltable {layout_attrs} frame="all">')
            xml_lines.append(f'      <tgroup cols="{element.colCount}">')
            xml_lines.append('        <tbody>')
            
            for row in element.rows:
                xml_lines.append('          <row>')
                for cell in row.cells:
                    # Cell layout? DocBook cells don't usually have x/y, but we can put it on entry if needed.
                    # For now, we just structure it.
                    colspan = f' namest="col{cell.colIndex}" nameend="col{cell.colIndex + cell.colSpan - 1}"' if cell.colSpan > 1 else ""
                    # DocBook spans are harder (namest/nameend), skipping complex span logic for now or assuming 1:1
                    
                    xml_lines.append(f'            <entry{colspan}>')
                    for item in cell.content:
                        if isinstance(item, TextObject):
                            safe_text = item.rawText.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                            xml_lines.append(f'              <para>{safe_text}</para>')
                    xml_lines.append('            </entry>')
                xml_lines.append('          </row>')
            
            xml_lines.append('        </tbody>')
            xml_lines.append('      </tgroup>')
            xml_lines.append('    </informaltable>')
