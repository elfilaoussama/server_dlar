"""
DOCX utilities for merging multiple documents.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List
import os

from api.core import get_logger

logger = get_logger("dla.docx_utils")


def merge_docx_files(docx_paths: List[str], output_path: str) -> str:
    """
    Merge multiple DOCX files into a single document.
    
    Args:
        docx_paths: List of paths to DOCX files (in order)
        output_path: Path for the merged output file
        
    Returns:
        Path to the merged document
    """
    if not docx_paths:
        raise ValueError("No DOCX files to merge")
    
    if len(docx_paths) == 1:
        # Only one file, just copy it
        import shutil
        shutil.copy(docx_paths[0], output_path)
        return output_path
    
    # Start with the first document
    merged_doc = Document(docx_paths[0])
    
    # Add page breaks and content from subsequent documents
    for i, docx_path in enumerate(docx_paths[1:], start=2):
        try:
            # Add page break before next document
            merged_doc.add_page_break()
            
            # Open the next document
            sub_doc = Document(docx_path)
            
            # Copy all elements from the sub document
            for element in sub_doc.element.body:
                # Skip sectPr (section properties) to avoid formatting issues
                if element.tag == qn('w:sectPr'):
                    continue
                merged_doc.element.body.append(element)
            
            logger.debug(f"Merged page {i} from {docx_path}")
            
        except Exception as e:
            logger.error(f"Failed to merge {docx_path}: {e}")
            continue
    
    # Save the merged document
    merged_doc.save(output_path)
    logger.info(f"Merged {len(docx_paths)} DOCX files into {output_path}")
    
    return output_path


def add_page_break(doc: Document):
    """Add a page break to the document."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run._r.append(OxmlElement('w:br'))
    run._r[-1].set(qn('w:type'), 'page')
