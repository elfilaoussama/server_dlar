from docx import Document
import os

def verify_docx(docx_path):
    if not os.path.exists(docx_path):
        print(f"Error: File not found at {docx_path}")
        return

    with open("verification_log.txt", "w") as f:
        try:
            doc = Document(docx_path)
            f.write(f"Successfully opened {docx_path}\n")
            
            # Check Page Dimensions
            section = doc.sections[0]
            f.write(f"Page Dimensions: {section.page_width.inches} x {section.page_height.inches} inches\n")
            
            # Count paragraphs
            num_paras = len(doc.paragraphs)
            f.write(f"Number of paragraphs: {num_paras}\n")
            
            # Count tables
            num_tables = len(doc.tables)
            f.write(f"Number of tables: {num_tables}\n")
            for i, table in enumerate(doc.tables):
                f.write(f"  Table {i}: {len(table.rows)} rows\n")
            
            # Check for frames (w:framePr)
            num_frames = 0
            num_images = 0
            text_content_len = 0
            
            for p in doc.paragraphs:
                # Check for frame properties (robust XML check)
                if p._p.pPr is not None and 'w:framePr' in p._p.pPr.xml:
                    num_frames += 1
                    # Extract Y coordinate
                    import xml.etree.ElementTree as ET
                    import re
                    match_y = re.search(r'w:y="(\d+)"', p._p.pPr.xml)
                    match_h = re.search(r'w:h="(\d+)"', p._p.pPr.xml)
                    
                    if match_y:
                        y_val = int(match_y.group(1))
                        h_val = int(match_h.group(1)) if match_h else 0
                        f.write(f"    Frame Y: {y_val} twips ({y_val/1440:.2f} in), H: {h_val} twips ({h_val/1440:.2f} in)\n")
                        
                        if y_val > 16000:
                            f.write("    ⚠️ WARNING: Frame Y is very large!\n")
                        if h_val > 16000:
                             f.write("    ⚠️ WARNING: Frame H is very large!\n")
                
                # Check for images (blip)
                if 'w:drawing' in p._p.xml or 'w:pict' in p._p.xml:
                    num_images += 1
                    
                text_content_len += len(p.text)

            f.write(f"Number of Framed Elements: {num_frames}\n")
            f.write(f"Number of Images (approx): {num_images}\n")
            f.write(f"Total Text Length: {text_content_len} characters\n")
            
            if num_frames > 0:
                f.write("✅ Layout Preservation: Frames detected.\n")
            else:
                f.write("❌ Layout Preservation: No frames found (Layout might be linear).\n")
                
            if text_content_len > 100:
                 f.write("✅ Content Extraction: Significant text content found.\n")
            else:
                 f.write("⚠️ Content Extraction: Low text content.\n")

        except Exception as e:
            f.write(f"Failed to inspect DOCX: {e}\n")

if __name__ == "__main__":
    # Path to the output file from the last run
    output_path = r"output\docx\reconstructed_document.docx"
    verify_docx(output_path)
